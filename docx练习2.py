#coding=utf-8
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn



path = os.getcwd()
test = Document()

#f = open('path/第一篇.docx','a+')这样错误
dizhi = path + '/' + '第一篇' + '.docx'

f = open(dizhi,'a+')
e = open('../nihao.txt','a+')

#f = open('aaa','a+')   带引号是字段不引用前面的变量，不带引号才引用

p = test.add_paragraph(u'第一段文字！')
run = p.add_run(u'\n24号字体！')

run.font.size = Pt(24)
run = p.add_run(u'\n中文字体')

run.font.name=u'楷体'

run.font.size = Pt(72)

r = run._element

r.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')




test.save(u'dizhi.docx')
#test.save(u'dizhi')就不行
