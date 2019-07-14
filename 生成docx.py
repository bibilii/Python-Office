import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_UNDERLINE
from docx.enum.dml import MSO_THEME_COLOR
from docx.shared import Inches,Pt,RGBColor

path = os.getcwd()
filename = 'test.docx'
file = path + '/' + filename
document = Document(file)

#添加文本
document.add_paragraph('这是第1次操作doc')

#上行间距
paragraph = document.add_paragraph('上行间距！')
paragraph_format = paragraph.paragraph_format
paragraph_format.space_before = Pt(40)  #18

#调整文本位置格式为居中：
paragraph = document.add_paragraph('文本居中')
paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

#调整文本位置格式为居左：
paragraph = document.add_paragraph('文本居左')
paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

#调整文本位置格式为居右：
paragraph = document.add_paragraph('文本居右')
paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

#调整左缩进0.3英寸
paragraph = document.add_paragraph('文本左缩进0.3英尺')
paragraph_format = paragraph.paragraph_format
paragraph_format.left_indent = Inches(0.3)

#下行间距
paragraph = document.add_paragraph('下行间距！')
paragraph_format = paragraph.paragraph_format
paragraph_format.space_after = Pt(12)


#首行缩进
paragraph = document.add_paragraph('首行缩进'*10)
paragraph_format = paragraph.paragraph_format
paragraph_format.first_line_indent = Inches(0.3)


#行距
paragraph = document.add_paragraph('行距'*20)
paragraph_format = paragraph.paragraph_format
paragraph_format.first_line_indent = Inches(0.3)
paragraph_format.line_spacing = Pt(18)

#分页格式
#紧跟上段：
paragraph = document.add_paragraph('紧跟上段'*100)
paragraph_format = paragraph.paragraph_format
paragraph_format.first_line_indent = Inches(0.3)
paragraph_format.keep_together

#若本页无法完全显示，另起一页：
paragraph = document.add_paragraph('若本页无法完全显示，另起一页'*200)
paragraph_format = paragraph.paragraph_format
paragraph_format.first_line_indent = Inches(0.3)
paragraph_format.keep_with_next

#强制另起一页：
paragraph = document.add_paragraph('强制另起一页'*200)
paragraph_format = paragraph.paragraph_format
paragraph_format.first_line_indent = Inches(0.3)
paragraph_format.page_break_before

#加粗    字体格式
p = document.add_paragraph()
run = p.add_run('字体格式：加粗')
run.font.bold = True

#斜体    字体格式
p = document.add_paragraph()
run = p.add_run('字体格式：斜体')
run.font.italic = True

#下划线    字体格式
p = document.add_paragraph()
run = p.add_run('字体格式：下划线')
run.font.underline = True

#WD_UNDERLINE 中有所有下划线格式
p = document.add_paragraph()
run = p.add_run('WD_UNDERLINE 中有所有下划线格式')
run.underline = WD_UNDERLINE.DOT_DASH
print(run.underline)

#字体颜色
test = document.add_paragraph('字体颜色').add_run('color')
font = test.font
font.color.rgb = RGBColor(0x42, 0x24 , 0xE9)

#调用预设颜色
test = document.add_paragraph('调用预设颜色').add_run('Color')
font = test.font
font.color.theme_color = MSO_THEME_COLOR.ACCENT_1

#添加图片
document.add_picture('部落二维码.png', width=Inches(1.25))
document.add_picture('公众号二维码.jpg', width=Inches(1.25))


document.save(file)
